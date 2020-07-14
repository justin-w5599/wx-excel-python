# -*- encoding: utf-8 -*-
'''
@File    :   pyWord.py
@Time    :   2020/07/01 11:05:27
@Author  :   Wang Junwen 
@Version :   1.0
@Contact :   junwen1938@163.com
'''

from docx import Document
from docx.shared import Inches
import os,re
rec=re.compile(r'IMG')
# print(dir(Document()))
path1=r"C:\Users\Administrator\Desktop\2020审核认定材料\初访台帐认定、审核材料\初访通过21案中9案没材料的"
os.chdir(path1)
namelist=os.listdir()
ln=[]
for n in namelist:
    if os.path.isdir(path1 + "\\" + n):
        ln.append(n)
stra=""
for l in ln:
    l1=os.listdir(path1 + "\\" + l)
    for n in l1:
        if not(re.match(rec,n)==None):
            stra=n
            doc=Document()
            doc.add_picture(path1 + "\\" + l + "\\" + stra, width=Inches(6.0))
            doc.save(path1 + "\\" + l + "\\" + l + r"五位一体包保表1.docx")    
        

    
# for x in ln:
#     os.system("del " + ".\\" + x + "\\" + "IMG*.jpg")        
# filename=os.listdir(r'./' + namelist[30])
# filename=filename.
# doc.add_picture(path1 + '\\' + namelist[0] + '\\' )